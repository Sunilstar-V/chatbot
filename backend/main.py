from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from typing import Any
from pymongo import MongoClient
import gridfs
import os
import traceback
from transformers import pipeline
from fastapi import Form
import PyPDF2
import docx
import pdfplumber
from io import BytesIO


# Load environment variables from .env file (if any)
load_dotenv()

app = FastAPI()

class Result(BaseModel):
    result: str | None 

# Update the origins based on your client-side app
origins = [
    "http://localhost",
    "http://localhost:8080",
    "http://localhost:3000"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

MONGODB_CONNECTION_STRING = os.getenv('MONGODB_CONNECTION_STRING')
DATABASE_NAME = 'chatbot_file_db'

# Create a MongoDB client instance
client = MongoClient(MONGODB_CONNECTION_STRING)
# Create (or reference) a MongoDB database
db = client[DATABASE_NAME]
# Create a GridFS file store instance in the specified MongoDB database
fs = gridfs.GridFS(db)

# Initialize the ChatGPT pipeline
# nlp = pipeline("text-generation", model="distilgpt2")
nlp = pipeline("text-generation", model="gpt2")


@app.post("/predict", response_model=Result)
async def predict(file: UploadFile = File(...), question: str = Form(...)) -> Any:
    try:
        if file is None:
            raise HTTPException(status_code=400, detail="File not uploaded.")
        if question is None or question.strip() == "":
            raise HTTPException(status_code=400, detail="No question provided.")

        # Save file content to MongoDB using GridFS
        # contents = await file.read()
        # fs.put(contents, filename=file.filename)

        # Your prediction code here...
        
        
        # nlp_input = f"Question: {question}\n\nFile Contents:\n{contents.decode()}"
        # result = nlp(nlp_input, max_length=100)[0]["generated_text"]

        # return {"result": result}
        
        bytes_len = len(await file.read())
        megabytes_len = bytes_len / (1024*1024)
        if megabytes_len > 100: # if file size > 100MB
            raise HTTPException(status_code=400, detail="File size exceeds limit (100MB).")

        allowed_extensions = ['.txt', '.csv', '.pdf', '.docx']
        file_extension = '.' + file.filename.split('.')[-1].lower()

        # Check if file has supported extension
        print(f"File extension: {file_extension}")
        if file_extension not in allowed_extensions:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        # Check if question was provided
        if question is None or question.strip() == "":
            raise HTTPException(status_code=400, detail="No question provided")

        # Read file contents
        contents = await file.read()
        file_data = await file.read()

        # Process file contents based on file type
        # if file_extension == '.txt' or file_extension == '.csv':
        #     contents = (await file.read()).decode('utf-8')

        # elif file_extension == '.docx':
        #     doc = docx.Document(BytesIO(await file.read()))
        #     contents = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

        # elif file_extension == '.pdf':
        #     with open("temporary.pdf", "wb") as out_file:
        #       out_file.write(file_data)

        #     # Open the pdf file with pdfplumber
        #     with pdfplumber.open("temporary.pdf") as pdf:
        #         contents = ''
        #         for page in pdf.pages:
        #             if page.extract_text() != None:
        #                 contents += ' ' + page.extract_text()
        
        your_file_text = "" # Initialize the variable
        if file_extension == '.txt' or file_extension == '.csv':
            your_file_text = contents.decode("utf-8")
        elif file_extension == '.docx':
            doc = docx.Document(BytesIO(contents))
            your_file_text = "\n".join([para.text for para in doc.paragraphs])
        elif file_extension == '.pdf':
            reader = PyPDF2.PdfReader(BytesIO(contents))
            your_file_text = "".join(page.extract_text() for page in reader.pages)

        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        nlp_input = f"File Contents:\n{your_file_text}\n\nQuestion: {question}"
        # Generate a response using the pipeline
        response = nlp(nlp_input, max_length=100)
        # Extract the generated text from the response
        generated_text = response[0]['generated_text']
        print(generated_text)
        return Result(result=generated_text)


    except HTTPException as http_err:
        print(traceback.format_exc())
        return {"result": str(http_err.detail)} # Pass the detail message of HTTPException
    except Exception as e:
        print(traceback.format_exc())
        return {"result": f"An error occurred: {str(e)}"}



# from dotenv import load_dotenv
# from fastapi import FastAPI, File, UploadFile
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel
# from transformers import pipeline
# from fastapi import HTTPException
# from pymongo import MongoClient
# import gridfs
# from typing import Optional
# from typing import Any

# # Load environment variables from .env file (if any)
# load_dotenv()

# app = FastAPI()

# class Response(BaseModel):
#     result: str | None  

# origins = [
#     "http://localhost",
#     "http://localhost:8080",
#     "http://localhost:3000"
# ]

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=origins,
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )
# #establishing connection with mongoDB to store files
# client = MongoClient('mongodb://localhost:27017/')
# db = client['chatbot_file_db']
# fs = gridfs.GridFS(db)

# def store_files(filename: str, file_content: bytes):
#     try:
#         fs.put(file_content, filename = filename)
#     except Exception as e:
#         print("An exception occurred while storing the file: ", e)
#         raise Exception("An exception occurred while storing the file: ") from e

# nlp = pipeline("text-generation", model="distilgpt2")

# @app.post("/predict", response_model = Response)
# # def predict() -> Any:

# async def predict(file: UploadFile = File(...), question: Optional[str] = None) -> Any:
    
#     print(f'Received file: {file.filename if file else None}')
#     print(f'Received question: {question}')
    
#     # Check if file was sent
#     if file is None:
#         raise HTTPException(status_code=400, detail="File not uploaded")

#     allowed_extensions = ['.txt', '.csv', '.pdf', '.docx']
#     file_extension = file.filename.split('.')[-1].lower()

#     # Check if file has supported extension
#     if file_extension not in allowed_extensions:
#         raise HTTPException(status_code=400, detail="Unsupported file type")

#     # Check if question was provided
#     if question is None or question.strip() == "":
#         raise HTTPException(status_code=400, detail="No question provided")

#     contents = await file.read()  # Moved line up
#     input_text = f"Question: {question}\n\nFile Contents:\n{contents.decode()}"

#     # File-specific processing replaced with common file handling for simplicity
#     result = None
#     if file_extension == 'txt':
#         # Perform processing specific to text files
#         result = "Processing text file"
#     elif file_extension == 'csv':
#         # Perform processing specific to CSV files
#         result = "Processing CSV file"
#     elif file_extension == 'docx':
#         # Perform processing specific to Word documents
#         result = "Processing Word document"
#     elif file_extension == 'pdf':
#         # Perform processing specific to PDF files
#         result = "Processing PDF file"
#     else:
#         result = f"Unsupported file type: .{file_extension}"

#     response = nlp(input_text, max_length=100)[0]["generated_text"]
#     return {"result": response or result}